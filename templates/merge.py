#!/usr/bin/env python3
# templates/merge.py
#
# Usage (run from project root):
#   python3 templates/merge.py \
#     --lab       "1" \
#     --title     "Разработка базы данных. Основы работы в MySQL" \
#     --subject   "Инжиниринг и управление данными" \
#     --group     "ИТИм-25" \
#     --student   "Скиндер И.П." \
#     --teacher   "доц. Романюк В.В." \
#     --city      "Донецк" \
#     --year      "2026" \
#     --content   "lab-01/_content_tmp.docx" \
#     --output    "lab-01/Отчет.docx"

import argparse, os, sys
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docxcompose.composer import Composer
from docx.enum.text import WD_LINE_SPACING
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Arguments ─────────────────────────────────────────────────────────────────

parser = argparse.ArgumentParser()
parser.add_argument('--lab',     required=True)
parser.add_argument('--title',   required=True)
parser.add_argument('--subject', required=True)
parser.add_argument('--group',   required=True)
parser.add_argument('--student', required=True)
parser.add_argument('--teacher', required=True)
parser.add_argument('--city',    default="Донецк")
parser.add_argument('--year',    required=True)
parser.add_argument('--content', required=True)
parser.add_argument('--output',  required=True)
args = parser.parse_args()

templates_dir  = os.path.dirname(__file__)
reference_path = os.path.join(templates_dir, 'reference.docx')

if not os.path.exists(reference_path):
    print(f'ERROR: {reference_path} not found'); sys.exit(1)
if not os.path.exists(args.content):
    print(f'ERROR: {args.content} not found'); sys.exit(1)

# == Step 1: Load reference.docx as the base document =========================
#
# KEY DECISION: we use reference.docx (not lab-template.docx) as the base.
# docxcompose keeps the BASE document's styles when there's a conflict.
# By starting from reference.docx, our carefully defined styles always win.
#
# lab-template.docx is only used as a visual reference — we recreate its
# content as paragraphs in code below, not by loading the file itself.

doc = Document(reference_path)

# == Step 2: Clear the placeholder paragraph from reference.docx ===============
#
# reference.docx has one placeholder paragraph ("Этот файл является шаблоном...").
# We remove it to get a blank document.
#
# Styles live in styles.xml — a completely separate part of the zip.
# Removing paragraphs only touches document.xml (the body content).
# The styles are untouched.

for para in doc.paragraphs:
    p = para._element           # access the raw XML element
    p.getparent().remove(p)     # remove it from its parent <w:body>

# == Step 3: Helper function for adding title page paragraphs ==================
#
# Every title page line is a Normal paragraph — we just vary:
#   - alignment (CENTER, RIGHT, or LEFT/inherited)
#   - font size (13pt for the university header, 14pt for everything else)
#   - the text content

def add_line(text, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=14, single_spacing=False):
    para = doc.add_paragraph()
    para.alignment = align
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)

    if single_spacing:
        # Override the 1.5 inherited from Normal style
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    return para

def add_blank():
    # A blank line is just a paragraph with no text
    add_line("")

# == Step 4: Build title page content =========================================
#
# Structure mirrors lab-template.docx exactly (verified by inspecting
# all 29 paragraphs including blank spacing lines):
#
# [0]  university ministry line   CENTER 14pt
# [1]  university name            CENTER 13pt
# [2-5] blank lines
# [6]  department                 RIGHT  14pt
# [7-10] blank lines
# [11] lab number                 CENTER 14pt
# [12] subject line               CENTER 14pt
# [13] topic line                 CENTER 14pt
# [14-17] blank lines
# [18] "Выполнил:"                RIGHT   14pt
# [19] group                      RIGHT   14pt
# [20] student name               RIGHT   14pt
# [21] blank
# [22] "Проверил:"                RIGHT   14pt
# [23] teacher name               RIGHT   14pt
# [24-27] blank lines
# [28] city + year                CENTER 14pt

add_line(
    "МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ "
    "федеральное государственное бюджетное образовательное учреждение "
    "высшего образования",
    size_pt=11,
    single_spacing=True
)
add_line('"ДОНЕЦКИЙ НАЦИОНАЛЬНЫЙ ТЕХНИЧЕСКИЙ УНИВЕРСИТЕТ"', size_pt=13)
add_blank(); add_blank(); add_blank(); add_blank()

add_line("Кафедра экономической кибернетики", align=WD_ALIGN_PARAGRAPH.RIGHT)
add_blank(); add_blank(); add_blank(); add_blank()

add_line(f"Лабораторная работа №{args.lab}")
add_line(f"по дисциплине «{args.subject}»")
add_line(f"на тему: «{args.title}»")
add_blank(); add_blank(); add_blank(); add_blank()

add_line("Выполнил:",          align=WD_ALIGN_PARAGRAPH.RIGHT)
add_line(f"ст. гр. {args.group}", align=WD_ALIGN_PARAGRAPH.RIGHT)
add_line(args.student,         align=WD_ALIGN_PARAGRAPH.RIGHT)
add_blank()

add_line("Проверил:",          align=WD_ALIGN_PARAGRAPH.RIGHT)
add_line(args.teacher,         align=WD_ALIGN_PARAGRAPH.RIGHT)
add_blank(); add_blank(); add_blank(); add_blank()

add_line(f"{args.city}, {args.year}")

# == Step 5: Page break before content ========================================
#
# A hard page break in docx XML is <w:br w:type="page"/> inside a run.
# We append it as a paragraph so content starts on a fresh page.

pb = OxmlElement('w:p')
r  = OxmlElement('w:r')
br = OxmlElement('w:br')
br.set(qn('w:type'), 'page')
r.append(br); pb.append(r)
doc.element.body.append(pb)

# == Step 6: Append content with docxcompose ===================================
#
# Composer handles numbering.xml, relationships, and image parts correctly.
# It remaps IDs from the content document to avoid conflicts with the base.

composer = Composer(doc)
composer.append(Document(args.content))

# == Step 7: Save ==============================================================

os.makedirs(os.path.dirname(os.path.abspath(args.output)), exist_ok=True)
composer.save(args.output)
print(f'Saved: {args.output}')
